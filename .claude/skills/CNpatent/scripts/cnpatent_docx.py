r"""
CNpatent DOCX helpers — single source of truth for python-docx operations
on the 交底书模板.

Why this file exists:
    Without it, every Writer / integration agent re-derives the same helper
    functions (clear_placeholders, get_section_anchors, insert_before, ...)
    from references/docx-patterns.md, which (a) wastes context tokens on
    every invocation and (b) risks divergence if one copy is patched but
    others are not. Bundling the helpers here lets the orchestrator simply
    `import` them.

Usage from Phase 3 (orchestrator):

    import sys
    from pathlib import Path

    skill_root = Path.cwd() / '.claude' / 'skills' / 'CNpatent'
    if not skill_root.exists():
        skill_root = Path.home() / '.claude' / 'skills' / 'CNpatent'
    sys.path.insert(0, str(skill_root / 'scripts'))

    from cnpatent_docx import (
        load_template, clear_placeholders, get_section_anchors,
        insert_before, append_at_end, add_formula, add_caption,
        para_replace, verify_docx,
    )

    doc = load_template()
    clear_placeholders(doc)
    anchors = get_section_anchors(doc)
    insert_before(anchors['一'], doc, '一种...方法')
    add_formula(doc, r'$\alpha = \arg\max_i f(x_i) \quad （1）$')
    doc.save(output_path)
    verify_docx(output_path)

All paragraph styles are handled by the template's styles.xml — do NOT
manually override `font.name` / `font.size` in client code. The whole
point of using a template is that the styles already encode the
correct font, size, and spacing required by CNIPA submission norms.

NOTE: this module-level docstring uses a raw-string prefix so that LaTeX
backslashes in the example above (\alpha, \arg, \max, \quad) don't trigger
Python SyntaxWarning during import.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


# ──────────────────────────────────────────────────────────────────
# Template loading
# ──────────────────────────────────────────────────────────────────

def find_template_path() -> Path:
    """Locate 交底书模板.docx in either project-level or global skill install."""
    candidates = [
        Path.cwd() / '.claude' / 'skills' / 'CNpatent' / 'assets' / '交底书模板.docx',
        Path.home() / '.claude' / 'skills' / 'CNpatent' / 'assets' / '交底书模板.docx',
    ]
    template = next((p for p in candidates if p.exists()), None)
    if template is None:
        raise FileNotFoundError(
            '找不到 交底书模板.docx，请确认 CNpatent skill 安装正确 '
            '(.claude/skills/CNpatent/assets/ 或 ~/.claude/skills/CNpatent/assets/)'
        )
    return template


def load_template() -> Document:
    """Load the bundled 交底书模板 and return a python-docx Document."""
    return Document(str(find_template_path()))


# ──────────────────────────────────────────────────────────────────
# Section anchor lookup
# ──────────────────────────────────────────────────────────────────

def find_sections(doc):
    """Return ordered list of (idx, level, text, paragraph) for all heading
    paragraphs (Heading 1 / Heading 2). Used by get_section_anchors().
    """
    sections = []
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ''
        if style_name in ('Heading 1', 'Heading 2'):
            sections.append((i, style_name, p.text.strip(), p))
    return sections


def get_section_anchors(doc):
    """Return {section_key: next_heading_paragraph} so callers can do
    `insert_before(anchor, doc, text)` to write content under each section.

    Section keys:
        '一', '二', '三', '四', '五', '六'
        '发明目的', '技术解决方案', '技术效果'  (children of 四)

    The last section ('六') has next=None — clients should append_at_end()
    instead of insert_before() for that section.
    """
    sections = find_sections(doc)
    keys = []
    for _, _, text, _ in sections:
        if text.startswith('一、'):
            keys.append('一')
        elif text.startswith('二、'):
            keys.append('二')
        elif text.startswith('三、'):
            keys.append('三')
        elif text.startswith('四、'):
            keys.append('四')
        elif text == '发明目的':
            keys.append('发明目的')
        elif text == '技术解决方案':
            keys.append('技术解决方案')
        elif text == '技术效果' or text.startswith('3、技术效果'):
            keys.append('技术效果')
        elif text.startswith('五、'):
            keys.append('五')
        elif text.startswith('六、'):
            keys.append('六')
        else:
            keys.append(None)

    next_anchor = {}
    for i, key in enumerate(keys):
        if key is None:
            continue
        if i + 1 < len(sections):
            next_anchor[key] = sections[i + 1][3]
        else:
            next_anchor[key] = None
    return next_anchor


# ──────────────────────────────────────────────────────────────────
# Placeholder cleanup
# ──────────────────────────────────────────────────────────────────

def clear_placeholders(doc):
    """Delete the template's pre-set empty / Caption / 图片 placeholder
    paragraphs.

    Preserves:
        - All Heading 1 / Heading 2 paragraphs (the section skeleton)
        - Any paragraph carrying a sectPr (section break — deleting these
          would corrupt page layout)
    """
    to_delete = []
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ''
        if style_name in ('Heading 1', 'Heading 2'):
            continue
        ppr = p._element.find(qn('w:pPr'))
        if ppr is not None and ppr.find(qn('w:sectPr')) is not None:
            continue
        if not p.text.strip() or style_name in ('Caption', '图片'):
            to_delete.append(i)
    for idx in sorted(to_delete, reverse=True):
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)
    return doc


# ──────────────────────────────────────────────────────────────────
# Content insertion
# ──────────────────────────────────────────────────────────────────

def add_body(doc, text, style='Normal'):
    """Append a paragraph using the named template style."""
    para = doc.add_paragraph(text)
    para.style = doc.styles[style]
    return para


def insert_before(anchor, doc, text, style='Normal'):
    """Insert a new paragraph immediately before `anchor`, inheriting the
    named template style. Used to write content under each numbered section.
    """
    para = add_body(doc, text, style)
    anchor._element.addprevious(para._element)
    return para


def append_at_end(doc, text, style='Normal'):
    """Append at document end. Used for 六、具体实施方式 (the last section,
    which has no following anchor)."""
    return add_body(doc, text, style)


def add_formula(doc, latex_text, anchor=None):
    """Add a plain-text LaTeX formula paragraph using the template's `公式` style.

    Args:
        doc: Document object
        latex_text: LaTeX string. **Must use Python raw literal** `r'...'` —
            otherwise `\\a` / `\\b` / `\\f` get interpreted as control
            characters (BEL / BS / FF), which are not valid XML and will
            crash the docx serialiser.
        anchor: If provided, insert before anchor; otherwise append to end.

    Examples:
        add_formula(doc, r'$\\alpha = \\arg\\max_i f(x_i) \\quad （1）$')
        add_formula(doc, r'$E = mc^2$', anchor=next_heading)
    """
    if anchor is not None:
        return insert_before(anchor, doc, latex_text, style='公式')
    return append_at_end(doc, latex_text, style='公式')


def add_caption(doc, text, anchor=None):
    """Add a figure caption paragraph using the template's `Caption` style."""
    if anchor is not None:
        return insert_before(anchor, doc, text, style='Caption')
    return append_at_end(doc, text, style='Caption')


# ──────────────────────────────────────────────────────────────────
# Post-modification: paragraph-level replace (handles run splitting)
# ──────────────────────────────────────────────────────────────────

def para_replace(paragraph, old, new):
    """Replace text within a paragraph, working around python-docx run splitting.

    Word internally splits paragraph text across multiple Run objects
    (e.g. "图6" may be stored as Run[0]="图" + Run[1]="6"), so searching
    a single Run for "图6" can fail. This function operates on the merged
    paragraph text and rewrites everything into the first Run, inheriting
    that Run's formatting.

    Returns True if a replacement was made, False otherwise.
    """
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    runs = paragraph.runs
    if not runs:
        return False
    runs[0].text = new_text
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    return True


# ──────────────────────────────────────────────────────────────────
# Post-write verification (Phase 4)
# ──────────────────────────────────────────────────────────────────

REQUIRED_HEADINGS = [
    '一、发明/实用新型名称',
    '二、所属技术领域',
    '三、现有技术（背景技术）',
    '四、发明内容',
    '发明目的',
    '技术解决方案',
    '技术效果',
    '五、附图及附图的简单说明',
    '六、具体实施方式',
]


def _has_any_formula(text):
    """Does this paragraph contain ANY LaTeX-like content (display or inline)?

    Used to enforce 'no formulas in 四、发明内容' — the legal-claim-drafting
    section must be pure prose, with not even inline `$x$` references.
    """
    return '$$' in text or bool(re.search(r'\$[^$\n]+\$', text))


def _is_display_formula_para(text):
    """Is this paragraph a *display* formula (its own line, dominated by
    the equation), as opposed to a prose paragraph that references LaTeX
    inline (e.g. '其中 $f(x_i)$ 为适应度值')?

    Heuristic: a display formula paragraph starts with `$` (after optional
    whitespace). Inline references start with Chinese/English prose.

    Display formulas should use the 公式 paragraph style; inline references
    should remain Normal (they're prose, not equations).
    """
    stripped = text.strip()
    return stripped.startswith('$')


def verify_docx(doc_path):
    """Run all post-write assertions on a generated 交底书 docx.

    Raises AssertionError on any violation:
        1. All 9 template headings present (6 H1 + 3 H2)
        2. No "权利要求书" leakage anywhere in the document
        3. Figure numbering 图1..图N in 五、附图说明 is contiguous
        4. Body text references don't exceed the max figure number
        5a. NO LaTeX of any kind (display or inline) outside 六、
            (because 发明内容 must be pure prose for legal claim drafting)
        5b. Display-formula paragraphs inside 六、 must use the 公式 style
            (inline `$x$` references in prose may stay Normal)
        6. Numbered items use full-width brackets （N） — no half-width (N)
    """
    doc_check = Document(str(doc_path))
    texts = [p.text for p in doc_check.paragraphs]

    # 1. Heading completeness
    heading_texts = [
        p.text.strip() for p in doc_check.paragraphs
        if p.style and p.style.name in ('Heading 1', 'Heading 2')
    ]
    for h in REQUIRED_HEADINGS:
        if not any(h in t for t in heading_texts):
            raise AssertionError(f'缺少标题段: {h}')

    # 2. No 权利要求书
    if any('权利要求书' in t for t in texts):
        raise AssertionError('不应包含权利要求书相关段落')

    # 3. Figure numbering contiguous in 五、附图说明.
    # We only consider paragraphs where the FIRST token is `图N` (i.e.
    # caption lines), not body references like "如图3所示".
    fig_nums = []
    for t in texts:
        m = re.match(r'^图\s*(\d+)', t.strip())
        if m:
            fig_nums.append(int(m.group(1)))
    if fig_nums:
        unique_sorted = sorted(set(fig_nums))
        if fig_nums != unique_sorted:
            raise AssertionError(f'附图编号不连续或有重复: {fig_nums}')

    # 4. Body references within range
    max_fig = max(fig_nums) if fig_nums else 0
    for t in texts:
        for m in re.finditer(r'图\s*(\d+)', t):
            ref_num = int(m.group(1))
            if ref_num > max_fig:
                raise AssertionError(f'引用了不存在的图{ref_num}（最大为图{max_fig}）')

    # 5. Formula constraints
    in_section_six = False
    for p in doc_check.paragraphs:
        text = p.text.strip()
        style = p.style.name if p.style else ''
        if style == 'Heading 1' and text.startswith('六、'):
            in_section_six = True
            continue

        # 5a. No LaTeX of any kind outside 六、 — even inline `$x$`
        if not in_section_six and style not in ('Heading 1', 'Heading 2'):
            if _has_any_formula(text):
                raise AssertionError(f'公式（含内联引用）出现在六节之外: {text[:50]}')

        # 5b. Display-formula paragraphs inside 六、 must use 公式 style
        if in_section_six and _is_display_formula_para(text):
            if style != '公式':
                raise AssertionError(f'公式段未使用 公式 样式 (实际: {style}): {text[:50]}')

    # 6. Numbered items use full-width brackets
    for t in texts:
        if re.search(r'^\(\d+\)', t.strip()):
            raise AssertionError(f'发现半角编号 (N)，应改为全角（N）: {t[:50]}')

    return True
